#!/usr/bin/env python3
"""
excel_to_word_findings.py
==========================

Convert a cybersecurity "Follow-up Plan" Excel workbook into a formal Word
report. The SAME data-extraction logic (locating the sheet/table, mapping
columns, validating values, grouping by section) is shared across THREE
output FORMATS, selected via --format:

    - "portrait-detail" (default): A4 portrait. Renders ONE 2-column Word
      table PER FINDING (General Control Review, Vulnerability Scanning,
      Web/API Penetration Testing, Source Code Review, etc), matching the
      detailed per-finding write-up layout used elsewhere in the report.

    - "landscape-detail": A4 landscape. Renders ONE summary Word table PER
      SECTION (e.g. "9.1 General Control Review", "9.2 Vulnerability
      Scanning", ...), with one row per finding and columns for #, Findings,
      Affected, Risk Description, Risk Level, Impact / Likelihood,
      Recommendation, and Rectification Status.

    - "veri-summary-by-section": A4 portrait. Renders ONE small summary
      table PER SECTION, counting the number of findings by Risk Level x
      Rectification (verification) Status - i.e. a "Verification Summary
      per Section" table, matching the sample layout provided by the user.

-------------------------------------------------------------------------
HOW IT LOCATES THE DATA
-------------------------------------------------------------------------
1. Opens the worksheet named "Follow-up Items" (falls back to the 3rd sheet
   in the workbook if that exact name isn't found, with a warning).
2. Scans every cell in the worksheet and finds the bounding box of ALL cells
   that have at least one visible border side set. This bounding box is
   assumed to be the findings table (borders are the only reliable signal,
   per the source workbook's convention, since the table doesn't always
   start at A1).
3. The first row of that bounding box is treated as the header row. Column
   purposes are identified by matching header text (case-insensitive,
   whitespace-normalized) against known labels:
       - 1st column (leftmost)      -> Finding ID (e.g. "W1", "MP1")
       - "Finding"                  -> Finding title
       - "Affected*"                -> Affected asset / URL / endpoint
       - "Risk Level"                -> Risk Level
       - "Impact" (optional)         -> Impact
       - "Likelihood" (optional)     -> Likelihood
       - "Risk Description"          -> Risk Description
       - "Recommend*"/"Safeguard*"   -> Recommended Safeguards
       - "Verification*"             -> one or more verification columns
                                         (rightmost non-empty per row wins)
   Any other column immediately to the left of the first "Verification*"
   column is treated as the (unused-in-output) Vendor Response column.
4. Rows where at least the first 6 columns of the table are merged into one
   cell are treated as SECTION HEADER rows (e.g. "General Control Review",
   "Vulnerability Scanning", "Web Penetration Testing") and are rendered as
   Word Heading 2 paragraphs instead of finding tables. The merge does not
   need to span the entire table width - only the first 6 columns (a merge
   spanning the FULL table width also still counts, since it necessarily
   covers at least the first 6 columns too).
5. Every other non-blank row within the bounding box is treated as a single
   finding and validated + rendered as its own 2-column Word table.

-------------------------------------------------------------------------
VALIDATION
-------------------------------------------------------------------------
The script warns (stderr) - but does NOT fail - on unexpected values for:
    - Risk Level     : Critical / High / Medium / Low / OFI
    - Impact         : Critical / High / Medium / Low / Very Low
    - Likelihood     : High / Medium / Low / Very Low
    - Verification   : the cell value must BEGIN WITH (case-insensitive) one
                        of: Completed / Partially Completed / Incomplete /
                        Scheduled / Accepted. If it does, the output value
                        is NORMALIZED to just that canonical label (the rest
                        of the cell's text, e.g. "Completed. Vendor added a
                        strict allow-list." -> "Completed", is dropped from
                        the Word table). If no match is found, a warning is
                        raised and the ORIGINAL (un-normalized) text is kept
                        in the output so nothing is silently lost.
All warnings are also collected and summarized at the end of the run.

-------------------------------------------------------------------------
WORD TABLE STYLING - "portrait-detail" format
-------------------------------------------------------------------------
    - All text (including section headings): Times New Roman, 12pt, black
      (Automatic) font color.
    - All cell content is top-aligned (vertically) and has no extra spacing
      between wrapped paragraphs within a cell.
    - First row of each finding table (Finding ID / Finding title):
      standard blue shading (#0070C0), white font, bold.
    - First column, all rows EXCEPT the first row: not bold.
    - Risk Level VALUE cell: shaded according to its risk level
      (Critical=#FF0000, High=#F4B083, Medium=#FFFF00, Low=#00FFFF,
      OFI=#92D050).
    - All other cells: no fill (transparent / white background).
    - Each finding's table is separated from the next by 2 blank lines
      (with no extra paragraph spacing added below them).

-------------------------------------------------------------------------
WORD TABLE STYLING - "landscape-detail" format
-------------------------------------------------------------------------
    - Page: A4, landscape orientation.
    - All text (including section headings and the "The following issues
      are identified:" intro paragraph): Times New Roman, 12pt, black
      (Automatic) font color.
    - Each section becomes its own bold heading, auto-numbered as
      "<section-number>.<n> <Section Title>" (e.g. "9.1 General Control
      Review"), followed by one summary table for that section's findings.
    - Table header row: standard blue shading (#0070C0), white font, bold,
      and "Repeat Header Rows" enabled (repeats on every page the table
      spans).
    - Table columns: #, Findings, Affected, Risk Description, Risk Level,
      Impact / Likelihood, Recommendation, Rectification Status as of
      <date> (date is derived per-section from the rightmost verification
      column that has any data in that section, using the same date
      extraction/normalization logic as the portrait format).
    - Risk Level value cells are colored using the EXACT SAME color scheme
      as the portrait format (Critical=#FF0000, High=#F4B083,
      Medium=#FFFF00, Low=#00FFFF, OFI=#92D050) - the coloring logic is
      shared/reused, not duplicated.
    - Data rows are NOT bold (only the header row and section headings are
      bold).

-------------------------------------------------------------------------
WORD TABLE STYLING - "veri-summary-by-section" format
-------------------------------------------------------------------------
    - Page: A4, portrait orientation.
    - All text: Times New Roman, 12pt, black (Automatic) font color.
    - ONE compact summary table PER SECTION (e.g. "Security Risk Assessment
      - General Control Review"), each with:
        - Row 1 (full-width, merged): section title, standard blue shading
          (#0070C0), white bold font.
        - Rows 2-3 (header block): "Risk Level" (merged vertically across
          both rows) | blank/"Total" | "Number of items by Rectification
          Status" (merged horizontally across the status columns) with the
          individual status column names below it. Same blue/white/bold
          styling as row 1. "Repeat Header Rows" enabled for rows 1-3.
        - One row per Risk Level (Critical [only if used anywhere in the
          workbook] / High / Medium / Low / OFI - the last 4 are ALWAYS
          shown, even with all-zero counts): risk level name is BOLD and
          shaded using the EXACT SAME color scheme as the other two formats
          (Critical=#FF0000, High=#F4B083, Medium=#FFFF00, Low=#00FFFF,
          OFI=#92D050); count cells are plain (not bold), no fill.
        - Final "Total" row: label not bold; all count cells bold; no fill.
    - Columns: Risk Level, Total, Completed, Partially Completed [only if
      used anywhere in the workbook], Incomplete, Scheduled, Accepted.
    - Whether the "Critical" row and "Partially Completed" column appear is
      decided ONCE, globally, across ALL findings in the workbook (not
      per-section) so every section's table has an identical, consistent
      shape.

-------------------------------------------------------------------------
USAGE
-------------------------------------------------------------------------
    python excel_to_word_findings.py input.xlsx
    python excel_to_word_findings.py input.xlsx output.docx
    python excel_to_word_findings.py input.xlsx --sheet "Follow-up Items"
    python excel_to_word_findings.py input.xlsx --format landscape-detail
    python excel_to_word_findings.py input.xlsx --format landscape-detail --section-number 9
    python excel_to_word_findings.py input.xlsx --format veri-summary-by-section
    python excel_to_word_findings.py input.xlsx --debug

--format accepts a string enum (not a boolean), so more formats can be
added later without breaking the CLI:
    - "portrait-detail"          (default) - one detailed table per finding, A4 portrait.
    - "landscape-detail"         - one summary table per section, A4 landscape.
    - "veri-summary-by-section"  - one verification-status-count table per section, A4 portrait.

--section-number sets the base report section number used to auto-number
section headings in "landscape-detail" (default: "9", producing "9.1",
"9.2", "9.3", ... in the order sections appear in the workbook). Ignored
for "portrait-detail" and "veri-summary-by-section".

Manual overrides (use if auto-detection of the table picks the wrong
region - e.g. if other bordered cells exist elsewhere on the sheet):
    --top-row N --left-col N --bottom-row N --right-col N
(1-based row/column numbers, as shown in Excel's row/column headers.)
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import openpyxl
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from openpyxl.worksheet.worksheet import Worksheet

# =============================================================================
# Constants / accepted vocabularies
# =============================================================================

DEFAULT_SHEET_NAME = "Follow-up Items"
DEFAULT_SHEET_INDEX_FALLBACK = 2  # zero-based -> 3rd sheet

# Output format enum (string values, NOT a boolean, so more formats can be
# added later without changing the CLI shape).
FORMAT_PORTRAIT_DETAIL = "portrait-detail"
FORMAT_LANDSCAPE_DETAIL = "landscape-detail"
FORMAT_VERI_SUMMARY_BY_SECTION = "veri-summary-by-section"
OUTPUT_FORMATS = [FORMAT_PORTRAIT_DETAIL, FORMAT_LANDSCAPE_DETAIL, FORMAT_VERI_SUMMARY_BY_SECTION]
DEFAULT_OUTPUT_FORMAT = FORMAT_PORTRAIT_DETAIL

DEFAULT_SECTION_NUMBER = "9"

RISK_LEVELS = {"critical", "high", "medium", "low", "ofi"}
IMPACT_LEVELS = {"critical", "high", "medium", "low", "very low"}
LIKELIHOOD_LEVELS = {"high", "medium", "low", "very low"}

# Canonical verification status labels (output values). The source cell text
# only needs to START WITH one of these (case-insensitive) - any trailing
# text (e.g. explanatory notes) is dropped from the Word output.
VERIFICATION_CANONICAL_LABELS = [
    "Completed",
    "Partially Completed",
    "Incomplete",
    "Scheduled",
    "Accepted",
]
# Sort longest-first so "Partially Completed" is checked before "Completed"
# would otherwise never incorrectly match as a prefix of it (they don't
# overlap as prefixes of one another, but this keeps matching unambiguous
# and future-proof if labels are edited).
_VERIFICATION_MATCH_ORDER = sorted(
    VERIFICATION_CANONICAL_LABELS, key=len, reverse=True
)

# ---- Styling ----
FONT_NAME = "Times New Roman"
FONT_SIZE = 12

HEADER_ROW_FILL = "0070C0"     # standard blue
HEADER_ROW_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)  # white

RISK_LEVEL_COLORS = {
    "critical": "FF0000",
    "high": "F4B083",
    "medium": "FFFF00",
    "low": "00FFFF",
    "ofi": "92D050",
}

# Display order/labels for risk levels, shared by every format that breaks
# counts down by risk level.
RISK_LEVEL_DISPLAY = {
    "critical": "Critical",
    "high": "High",
    "medium": "Medium",
    "low": "Low",
    "ofi": "OFI",
}

WARNINGS: list[str] = []


def warn(message: str) -> None:
    """Record + immediately print a validation warning."""
    WARNINGS.append(message)
    print(f"WARNING: {message}", file=sys.stderr)


# =============================================================================
# Step 1-2: locate worksheet + table bounding box (via borders)
# =============================================================================


def load_sheet(wb: openpyxl.Workbook, sheet_name: str) -> Worksheet:
    if sheet_name in wb.sheetnames:
        return wb[sheet_name]

    # case-insensitive match
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return wb[name]

    if len(wb.sheetnames) > DEFAULT_SHEET_INDEX_FALLBACK:
        fallback = wb.sheetnames[DEFAULT_SHEET_INDEX_FALLBACK]
        warn(
            f'Sheet named "{sheet_name}" not found. '
            f'Falling back to the 3rd sheet: "{fallback}".'
        )
        return wb[fallback]

    raise ValueError(
        f'Sheet named "{sheet_name}" not found, and the workbook has fewer '
        f"than 3 sheets to fall back to. Available sheets: {wb.sheetnames}"
    )


def _cell_has_border(cell) -> bool:
    b = cell.border
    for side in (b.top, b.bottom, b.left, b.right):
        if side is not None and side.style is not None:
            return True
    return False


def find_table_bounds(ws: Worksheet) -> tuple[int, int, int, int]:
    """Return (min_row, min_col, max_row, max_col) of all bordered cells."""
    min_row = min_col = max_row = max_col = None
    for row in ws.iter_rows():
        for cell in row:
            if _cell_has_border(cell):
                r, c = cell.row, cell.column
                min_row = r if min_row is None else min(min_row, r)
                max_row = r if max_row is None else max(max_row, r)
                min_col = c if min_col is None else min(min_col, c)
                max_col = c if max_col is None else max(max_col, c)

    if min_row is None:
        raise ValueError(
            "No bordered cells were found on the sheet - cannot locate the "
            "findings table. Use --top-row/--left-col/--bottom-row/--right-col "
            "to specify the table location manually."
        )
    return min_row, min_col, max_row, max_col


# =============================================================================
# Text cleaning helpers
# =============================================================================

_TAG_RE = re.compile(r"<[^>]+>")
_DATE_RE = re.compile(
    r"(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{2,4})"
)
_MONTHS = [
    "Jan", "Feb", "Mar", "Apr", "May", "Jun",
    "Jul", "Aug", "Sep", "Oct", "Nov", "Dec",
]


def clean_text(value) -> str:
    """Convert a raw cell value to plain text, stripping any HTML markup
    that may have been pasted into the cell (e.g. anchor tags), while
    PRESERVING any surrounding plain text and the cell's original line
    breaks.

    Handles both:
        - A cell whose ENTIRE value is just an anchor tag, e.g.
          '<a href="https://x">https://x</a>' -> "https://x"
        - A cell with an anchor EMBEDDED within a larger block of text,
          e.g. 'Update the packages.\\n\\nReference: <a href="...">...</a>'
          -> 'Update the packages.\\n\\nReference: https://...' (all
          surrounding text is preserved; only the <a> tag itself is
          replaced by its visible text, or its href if it has none).
    """
    if value is None:
        return ""
    text = str(value)
    if "<" in text and ">" in text:
        try:
            soup = BeautifulSoup(text, "html.parser")
            for a_tag in soup.find_all("a"):
                visible = a_tag.get_text(strip=True)
                href = a_tag.get("href")
                a_tag.replace_with(visible or href or "")
            # No separator: real line breaks already exist as literal "\n"
            # characters WITHIN the original text nodes. Using a separator
            # here would incorrectly insert extra line breaks at every tag
            # boundary (e.g. splitting "Reference: " from the URL that
            # replaced the <a> tag onto two different lines).
            plain = soup.get_text().strip()
            if plain:
                return plain
        except Exception:
            return _TAG_RE.sub("", text).strip()
    return text.strip()


def split_paragraphs(value) -> list[str]:
    """Clean a (possibly multi-line) cell value and split it into a list of
    non-empty paragraphs, preserving the author's paragraph breaks."""
    text = clean_text(value)
    if not text:
        return []
    parts = [p.strip() for p in text.replace("\r\n", "\n").split("\n")]
    return [p for p in parts if p]


def extract_date(header_text: str) -> Optional[str]:
    """Try to pull a date out of a verification column header such as
    'Verification by Pentastic on 29.08.2026' and format it as 'dd MMM yyyy'.
    Returns None if no recognizable date pattern is found."""
    m = _DATE_RE.search(header_text)
    if not m:
        return None
    d, mth, y = m.groups()
    try:
        d, mth = int(d), int(mth)
        y = int(y)
        if y < 100:
            y += 2000
        if not (1 <= mth <= 12 and 1 <= d <= 31):
            return None
        return f"{d:02d} {_MONTHS[mth - 1]} {y}"
    except (ValueError, IndexError):
        return None


def normalize(s) -> str:
    return re.sub(r"\s+", " ", str(s or "")).strip().lower()


def normalize_verification_status(value: str) -> tuple[str, bool]:
    """Match `value` against the canonical verification labels as a
    case-insensitive PREFIX match. Returns (output_value, matched):
        - If a canonical label matches as a prefix, returns
          (canonical_label, True) - trailing text (e.g. explanatory notes)
          is dropped.
        - If no canonical label matches, returns (original_value, False) so
          the original text is preserved in the output (nothing is silently
          lost) and the caller can raise a warning.
        - Empty input returns ("", True) - no warning for blank cells.
    """
    if not value:
        return "", True

    norm_value = normalize(value)
    for label in _VERIFICATION_MATCH_ORDER:
        if norm_value.startswith(label.lower()):
            return label, True

    return value, False


# =============================================================================
# Step 3: header mapping
# =============================================================================


@dataclass
class HeaderMap:
    id_col: int
    finding_col: Optional[int] = None
    affected_col: Optional[int] = None
    risk_level_col: Optional[int] = None
    impact_col: Optional[int] = None
    likelihood_col: Optional[int] = None
    risk_description_col: Optional[int] = None
    recommended_safeguards_col: Optional[int] = None
    vendor_response_col: Optional[int] = None
    # list of (column_index, raw_header_text), left-to-right
    verification_cols: list[tuple[int, str]] = field(default_factory=list)


def map_headers(ws: Worksheet, header_row: int, min_col: int, max_col: int) -> HeaderMap:
    hmap = HeaderMap(id_col=min_col)
    assigned: set[int] = {min_col}

    for col in range(min_col + 1, max_col + 1):
        raw = ws.cell(row=header_row, column=col).value
        text = clean_text(raw)
        norm = normalize(text)
        if not norm:
            continue

        if norm.startswith("verification"):
            hmap.verification_cols.append((col, text))
            assigned.add(col)
        elif "finding" in norm:
            hmap.finding_col = col
            assigned.add(col)
        elif "affected" in norm:
            hmap.affected_col = col
            assigned.add(col)
        elif "risk level" in norm or norm == "risk":
            hmap.risk_level_col = col
            assigned.add(col)
        elif "risk description" in norm or ("description" in norm and "risk" in norm):
            hmap.risk_description_col = col
            assigned.add(col)
        elif "likelihood" in norm:
            hmap.likelihood_col = col
            assigned.add(col)
        elif norm == "impact" or ("impact" in norm and "risk" not in norm):
            hmap.impact_col = col
            assigned.add(col)
        elif "recommend" in norm or "safeguard" in norm:
            hmap.recommended_safeguards_col = col
            assigned.add(col)
        # else: leave unassigned for now (candidate for vendor response)

    # Vendor Response: header text/name is inconsistent across projects.
    # It's positionally the column immediately before the first
    # Verification* column (and is not consumed by the final Word output,
    # but we identify it so it isn't mistaken for anything else).
    if hmap.verification_cols:
        first_verif_col = hmap.verification_cols[0][0]
        candidate = first_verif_col - 1
        if candidate >= min_col and candidate not in assigned:
            hmap.vendor_response_col = candidate
            assigned.add(candidate)
    else:
        warn(
            "No 'Verification...' column was detected in the header row. "
            "Rectification status will be left blank for all findings."
        )

    if hmap.finding_col is None:
        warn('Could not find a "Finding" column in the header row.')
    if hmap.risk_description_col is None:
        warn('Could not find a "Risk Description" column in the header row.')
    if hmap.risk_level_col is None:
        warn('Could not find a "Risk Level" column in the header row.')
    if hmap.recommended_safeguards_col is None:
        warn('Could not find a "Recommended Safeguards" column in the header row.')

    return hmap


# =============================================================================
# Section header row detection (rows with at least the first 6 columns
# merged, e.g. "Penetration Testing")
# =============================================================================

SECTION_HEADER_MIN_MERGED_COLS = 6


def section_title_for_row(ws: Worksheet, row: int, min_col: int, max_col: int) -> Optional[str]:
    width = max_col - min_col + 1
    # A table narrower than the usual minimum can never satisfy the "first 6
    # columns merged" rule, so fall back to requiring the full width instead.
    required_merged_cols = min(SECTION_HEADER_MIN_MERGED_COLS, width)

    for mc in ws.merged_cells.ranges:
        if mc.min_row == row and mc.max_row == row:
            covered = mc.max_col - mc.min_col + 1
            # Treat as a section header if the merge starts at (or right at)
            # the leftmost column of the table AND spans at least the first
            # 6 columns (rather than requiring the entire row to be merged).
            # A row merged across the FULL table width also satisfies this,
            # since it necessarily covers at least the first 6 columns too.
            if mc.min_col <= min_col + 1 and covered >= required_merged_cols:
                anchor = ws.cell(row=mc.min_row, column=mc.min_col).value
                return clean_text(anchor)
    return None


# =============================================================================
# Step 3-4: extract + validate findings
# =============================================================================


@dataclass
class Finding:
    row: int
    finding_id: str
    finding_title: str
    affected: str
    risk_level: str
    impact: str
    likelihood: str
    risk_description: list[str]
    recommended_safeguards: list[str]
    verification_status: str
    verification_date_label: str
    # Column index (1-based, openpyxl-style) of the verification column that
    # was actually used (last non-empty, left-to-right) for THIS finding, or
    # None if no verification column had a value. Used by the
    # "landscape-detail" format to derive a single, section-level
    # "Rectification Status as of <date>" column header.
    verification_col_used: Optional[int] = None


def get_val(ws: Worksheet, row: int, col: Optional[int]) -> str:
    if col is None:
        return ""
    return clean_text(ws.cell(row=row, column=col).value)


def validate_risk_level(value: str, finding_id: str, row: int) -> None:
    if value and normalize(value) not in RISK_LEVELS:
        warn(
            f'Row {row} (Finding "{finding_id}"): unexpected Risk Level '
            f'"{value}". Expected one of {sorted(RISK_LEVELS)}.'
        )


def validate_impact(value: str, finding_id: str, row: int) -> None:
    if value and normalize(value) not in IMPACT_LEVELS:
        warn(
            f'Row {row} (Finding "{finding_id}"): unexpected Impact '
            f'"{value}". Expected one of {sorted(IMPACT_LEVELS)}.'
        )


def validate_likelihood(value: str, finding_id: str, row: int) -> None:
    if value and normalize(value) not in LIKELIHOOD_LEVELS:
        warn(
            f'Row {row} (Finding "{finding_id}"): unexpected Likelihood '
            f'"{value}". Expected one of {sorted(LIKELIHOOD_LEVELS)}.'
        )


def extract_findings(
    ws: Worksheet, hmap: HeaderMap, header_row: int, min_col: int, max_col: int, max_row: int
) -> list[tuple[Optional[str], list[Finding]]]:
    """Returns a list of (section_title_or_None, [Finding, ...]) groups, in
    top-to-bottom order, matching the layout of the source sheet."""

    groups: list[tuple[Optional[str], list[Finding]]] = []
    current_section: Optional[str] = None
    current_findings: list[Finding] = []

    for row in range(header_row + 1, max_row + 1):
        section = section_title_for_row(ws, row, min_col, max_col)
        if section is not None:
            if current_findings or current_section is not None:
                groups.append((current_section, current_findings))
            current_section = section
            current_findings = []
            continue

        finding_id = get_val(ws, row, hmap.id_col)
        finding_title = get_val(ws, row, hmap.finding_col)

        # Skip fully blank rows
        row_values = [
            get_val(ws, row, c)
            for c in (
                hmap.id_col,
                hmap.finding_col,
                hmap.affected_col,
                hmap.risk_level_col,
                hmap.risk_description_col,
            )
        ]
        if not any(row_values):
            continue

        risk_level = get_val(ws, row, hmap.risk_level_col)
        impact = get_val(ws, row, hmap.impact_col)
        likelihood = get_val(ws, row, hmap.likelihood_col)
        affected = get_val(ws, row, hmap.affected_col)
        risk_description = split_paragraphs(ws.cell(row=row, column=hmap.risk_description_col).value) if hmap.risk_description_col else []
        recommended_safeguards = split_paragraphs(ws.cell(row=row, column=hmap.recommended_safeguards_col).value) if hmap.recommended_safeguards_col else []

        # Verification: take the LAST non-empty column, left-to-right.
        verification_raw = ""
        verification_header = ""
        verification_col_used: Optional[int] = None
        for col, header_text in hmap.verification_cols:
            val = get_val(ws, row, col)
            if val:
                verification_raw = val
                verification_header = header_text
                verification_col_used = col

        # Normalize to one of the 5 canonical labels (prefix match), dropping
        # any trailing explanatory text. Warn if nothing matches.
        verification_status, matched = normalize_verification_status(verification_raw)
        if not matched:
            warn(
                f'Row {row} (Finding "{finding_id}"): unexpected verification '
                f'status "{verification_raw}". Expected it to start with one '
                f"of {VERIFICATION_CANONICAL_LABELS} (case-insensitive)."
            )

        date_str = extract_date(verification_header) if verification_header else None
        if date_str:
            verification_date_label = f"Rectification status as of {date_str}"
        elif verification_header:
            verification_date_label = f"Rectification status ({verification_header})"
        else:
            verification_date_label = "Rectification status as of dd MMM yyyy"

        validate_risk_level(risk_level, finding_id, row)
        validate_impact(impact, finding_id, row)
        validate_likelihood(likelihood, finding_id, row)

        current_findings.append(
            Finding(
                row=row,
                finding_id=finding_id,
                finding_title=finding_title,
                affected=affected,
                risk_level=risk_level,
                impact=impact,
                likelihood=likelihood,
                risk_description=risk_description,
                recommended_safeguards=recommended_safeguards,
                verification_status=verification_status,
                verification_date_label=verification_date_label,
                verification_col_used=verification_col_used,
            )
        )

    if current_findings or current_section is not None:
        groups.append((current_section, current_findings))

    return groups


def group_verification_status_label(findings: list[Finding], hmap: HeaderMap) -> str:
    """Derive a SINGLE 'Rectification Status as of <date>' column header for
    an entire section's summary table (used by the "landscape-detail"
    format, where the rectification/verification date is a per-TABLE column
    header rather than a per-ROW label as in "portrait-detail").

    Convention (reusing the same date-extraction logic as the per-row
    label): among all verification columns actually used by ANY finding in
    this section, pick the RIGHTMOST one (i.e. the latest verification
    round that applies to this section) and extract its date the same way
    the portrait format does. Falls back to a generic placeholder if no
    verification data is available for this section.
    """
    if not hmap.verification_cols:
        return "Rectification Status as of dd MMM yyyy"

    used_cols = {f.verification_col_used for f in findings if f.verification_col_used is not None}
    if not used_cols:
        return "Rectification Status as of dd MMM yyyy"

    max_col = max(used_cols)
    header_text = next((h for c, h in hmap.verification_cols if c == max_col), "")

    date_str = extract_date(header_text) if header_text else None
    if date_str:
        return f"Rectification Status as of {date_str}"
    elif header_text:
        return f"Rectification Status ({header_text})"
    return "Rectification Status as of dd MMM yyyy"


# =============================================================================
# Shared Word-building helpers (used by ALL formats)
# =============================================================================

LABEL_COL_WIDTH = Cm(4.2)
VALUE_COL_WIDTH = Cm(12.8)


def _set_table_column_widths(table, widths: list) -> None:
    """Explicitly (re)write the table's <w:tblGrid> to match the given
    column widths (a list of python-docx Length objects, e.g. Cm(1.2)), and
    also set <w:tblW> to the sum of those widths.

    Why this is necessary: python-docx's `cell.width = ...` only sets the
    per-cell <w:tcW>. Some renderers use <w:tblGrid> (not the individual
    per-cell widths) to lay out a table's columns when the two disagree -
    which they do by default, since add_table() always creates a tblGrid
    with the page's default width divided evenly across all columns. If we
    only set cell.width without also correcting tblGrid, columns can render
    at the wrong (evenly-distributed) widths despite tcW being "correct" in
    the underlying XML. This function keeps both in sync so the intended
    widths are honored consistently.
    """
    tbl = table._tbl
    tblGrid = tbl.tblGrid
    for gc in list(tblGrid):
        tblGrid.remove(gc)
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w.twips))
        tblGrid.append(gridCol)

    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(sum(w.twips for w in widths)))

    # Keep every existing row's cells in sync with the corrected grid too.
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def _set_cell_fill(cell, hex_color: Optional[str]) -> None:
    """Set the cell background fill. Pass hex_color=None for 'no fill'
    (transparent / default white background)."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove any existing shading element first
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)

    if hex_color is None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "auto")
        tcPr.append(shd)
    else:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _set_cell_text(
    cell,
    paragraphs: list[str],
    bold: bool = False,
    font_color: Optional[RGBColor] = None,
) -> None:
    """Write one or more paragraphs of text into a table cell, applying the
    document-wide font (Times New Roman, 12pt) with no extra spacing
    between wrapped paragraphs."""
    if not paragraphs:
        paragraphs = [""]
    cell.text = ""  # clear default empty paragraph's placeholder run
    first = True
    for text in paragraphs:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        run.font.bold = bold
        if font_color is not None:
            run.font.color.rgb = font_color
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def _add_blank_separator_paragraph(document: Document):
    """Add a blank paragraph used purely as vertical whitespace between
    tables, with all paragraph spacing zeroed out so it doesn't introduce
    any *extra* space beyond the blank line itself."""
    p = document.add_paragraph("")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
    return p


def _set_repeat_header_row(row) -> None:
    """Mark a table row as a repeating header row (OOXML <w:tblHeader/>),
    so it repeats at the top of every page the table spans."""
    trPr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def _hard_set_style_font(style, name: str = FONT_NAME, size_pt: int = FONT_SIZE) -> None:
    """Force a paragraph style's run properties (rPr) to an EXPLICIT font
    name/size/color, fully overriding Word's built-in theme-based styling.

    This is needed because Word's default "Heading 1"/"Heading 2" styles
    (and similar) don't just set a plain font name/size/color - they
    reference the document THEME instead:
        <w:rFonts w:asciiTheme="majorHAnsi" w:eastAsiaTheme="majorEastAsia"
                  w:hAnsiTheme="majorHAnsi" w:cstheme="majorBidi"/>
        <w:color w:val="365F91" w:themeColor="accent1" w:themeShade="BF"/>

    Simply setting `style.font.name = ...` (as python-docx's high-level API
    does) only ADDS explicit w:ascii/w:hAnsi attributes - it does NOT
    remove the w:asciiTheme/w:hAnsiTheme/etc. attributes already present.
    When both explicit and theme font references coexist, some renderers
    (observed with LibreOffice) still preferred the THEME font over the
    explicit one, so headings kept rendering in the theme's sans-serif
    heading font (and theme accent color) instead of Times New Roman
    black/auto. This function removes ALL theme references and writes
    fully explicit values so there's no ambiguity for any renderer.
    """
    # Set the size color etc via the standard python-docx API first (keeps
    # everything consistent / handles version differences gracefully).
    style.font.name = name
    style.font.size = Pt(size_pt)

    rPr = style.element.get_or_add_rPr()

    # ---- Fonts: remove theme references, set every font slot explicitly ----
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rFonts.get(qn(theme_attr)) is not None:
            del rFonts.attrib[qn(theme_attr)]
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)

    # ---- Size: also fix complex-script size (szCs), which python-docx's
    # high-level API does NOT update, to keep it consistent with sz. ----
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(size_pt * 2))

    # ---- Color: force to "auto" (Word's "Automatic" color, renders as
    # black), removing any themeColor/themeTint/themeShade references. ----
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    for theme_attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        if color.get(qn(theme_attr)) is not None:
            del color.attrib[qn(theme_attr)]
    color.set(qn("w:val"), "auto")


def _apply_base_styles(document: Document) -> None:
    """Shared styling setup used by ALL output formats: fixes the OOXML
    <w:zoom> validation issue, and forces Times New Roman 12pt, black
    (Automatic) font color across the Normal style AND the Heading 1/2
    styles, so ALL text in the document (body text and headings alike)
    uses the same font/size/color - overriding Word's theme-based heading
    defaults (see _hard_set_style_font for why this requires more than
    just setting .font.name/.font.size)."""
    # python-docx's default template omits the required w:percent attribute
    # on <w:zoom>; add it so the resulting file passes strict OOXML validation.
    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    _hard_set_style_font(document.styles["Normal"], FONT_NAME, FONT_SIZE)
    style = document.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for heading_style_id in ("Heading 1", "Heading 2"):
        if heading_style_id in document.styles:
            _hard_set_style_font(document.styles[heading_style_id], FONT_NAME, FONT_SIZE)


# =============================================================================
# "portrait-detail" format: one detailed table per finding, A4 portrait
# =============================================================================


def add_finding_table(document: Document, finding: Finding) -> None:
    table = document.add_table(rows=10, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_column_widths(table, [LABEL_COL_WIDTH, VALUE_COL_WIDTH])

    rows_content = [
        (finding.finding_id, [finding.finding_title]),
        ("Risk Description", finding.risk_description),
        ("Risk Level", [finding.risk_level]),
        ("Impact/Likelihood", [f"{finding.impact} / {finding.likelihood}" if (finding.impact or finding.likelihood) else ""]),
        ("OWASP Top 10", [""]),
        ("Affected Asset", [finding.affected]),
        ("Evidence for the finding", [""]),
        ("Recommended Safeguards", finding.recommended_safeguards),
        ("Evidence for the remedial actions", [""]),
        (finding.verification_date_label, [finding.verification_status]),
    ]

    RISK_LEVEL_ROW_INDEX = 2

    for i, (label, value_paragraphs) in enumerate(rows_content):
        row = table.rows[i]
        label_cell, value_cell = row.cells[0], row.cells[1]

        if i == 0:
            # First row: standard blue shading, white bold font (both cells)
            _set_cell_text(label_cell, [label], bold=True, font_color=HEADER_ROW_FONT_COLOR)
            _set_cell_text(value_cell, value_paragraphs, bold=True, font_color=HEADER_ROW_FONT_COLOR)
            _set_cell_fill(label_cell, HEADER_ROW_FILL)
            _set_cell_fill(value_cell, HEADER_ROW_FILL)
        else:
            _set_cell_text(label_cell, [label], bold=False)
            _set_cell_text(value_cell, value_paragraphs, bold=False)
            _set_cell_fill(label_cell, None)  # no fill

            if i == RISK_LEVEL_ROW_INDEX:
                risk_hex = RISK_LEVEL_COLORS.get(normalize(finding.risk_level))
                _set_cell_fill(value_cell, risk_hex)  # None if unrecognized -> no fill
            else:
                _set_cell_fill(value_cell, None)  # no fill

        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def build_document(groups: list[tuple[Optional[str], list[Finding]]], title: str) -> Document:
    """Build the "portrait-detail" output: A4 portrait, one 2-column Word
    table PER FINDING."""
    document = Document()
    _apply_base_styles(document)

    document.add_heading(title, level=1)

    total = 0
    for section_title, findings in groups:
        if not findings:
            continue
        if section_title:
            document.add_heading(section_title, level=2)
        for finding in findings:
            add_finding_table(document, finding)
            total += 1
            # Separate each table with 2 empty lines (no extra paragraph
            # spacing added below them).
            _add_blank_separator_paragraph(document)
            _add_blank_separator_paragraph(document)

    if total == 0:
        warn("No findings were extracted - the output document will be empty of tables.")

    return document


# =============================================================================
# "landscape-detail" format: one summary table per section, A4 landscape
# =============================================================================

# Content width targeted for A4 landscape with the margins set in
# _setup_a4_landscape() below: 29.7cm - (1.5cm left + 1.5cm right) = 26.7cm.
LANDSCAPE_COLUMNS = [
    # (header_label, width)
    ("#", Cm(1.2)),
    ("Findings", Cm(3.5)),
    ("Affected", Cm(3.6)),
    ("Risk Description", Cm(5.3)),
    ("Risk Level", Cm(2.0)),
    ("Impact / Likelihood", Cm(2.5)),
    ("Recommendation", Cm(5.1)),
    # The last column's header is dynamic ("Rectification Status as of
    # <date>") and is filled in per-section at render time; the width is
    # fixed here.
    (None, Cm(3.5)),
]


def _setup_a4_landscape(document: Document) -> None:
    """Configure the document's first section as A4, landscape orientation."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def _lines_from_text(text: str) -> list[str]:
    """Split an already-cleaned string (e.g. Finding.affected, which may
    retain literal '\\n' line breaks copied from the source Excel cell,
    such as multiple IP addresses) into a list of non-empty lines."""
    if not text:
        return [""]
    parts = [p.strip() for p in text.replace("\r\n", "\n").split("\n")]
    parts = [p for p in parts if p]
    return parts or [""]


def add_section_summary_table(
    document: Document,
    findings: list[Finding],
    hmap: HeaderMap,
) -> None:
    """Render ONE section's findings as a single summary table (one row per
    finding), per the "landscape-detail" layout."""
    n_cols = len(LANDSCAPE_COLUMNS)
    column_widths = [width for _, width in LANDSCAPE_COLUMNS]
    table = document.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_column_widths(table, column_widths)

    rectification_header = group_verification_status_label(findings, hmap)

    # ---- Header row ----
    header_row = table.rows[0]
    for i, (label, width) in enumerate(LANDSCAPE_COLUMNS):
        cell = header_row.cells[i]
        cell.width = width
        header_text = label if label is not None else rectification_header
        _set_cell_text(cell, [header_text], bold=True, font_color=HEADER_ROW_FONT_COLOR)
        _set_cell_fill(cell, HEADER_ROW_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_repeat_header_row(header_row)

    # ---- Data rows (one per finding) ----
    for finding in findings:
        row = table.add_row()
        cells = row.cells
        for i, (_, width) in enumerate(LANDSCAPE_COLUMNS):
            cells[i].width = width

        values = [
            [finding.finding_id],
            [finding.finding_title],
            _lines_from_text(finding.affected),
            finding.risk_description or [""],
            [finding.risk_level],
            [f"{finding.impact} / {finding.likelihood}" if (finding.impact or finding.likelihood) else ""],
            finding.recommended_safeguards or [""],
            [finding.verification_status],
        ]

        RISK_LEVEL_COL_INDEX = 4

        for i, paragraphs in enumerate(values):
            cell = cells[i]
            _set_cell_text(cell, paragraphs, bold=False)
            if i == RISK_LEVEL_COL_INDEX:
                # Reuse the EXACT SAME risk-level color scheme/logic as the
                # portrait format (RISK_LEVEL_COLORS + _set_cell_fill).
                risk_hex = RISK_LEVEL_COLORS.get(normalize(finding.risk_level))
                _set_cell_fill(cell, risk_hex)
            else:
                _set_cell_fill(cell, None)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def build_landscape_document(
    groups: list[tuple[Optional[str], list[Finding]]],
    hmap: HeaderMap,
    title: str,
    section_base_number: str = DEFAULT_SECTION_NUMBER,
) -> Document:
    """Build the "landscape-detail" output: A4 landscape, one summary table
    PER SECTION (e.g. "9.1 General Control Review"), reusing the same
    Finding/HeaderMap data produced by extract_findings()."""
    document = Document()
    _apply_base_styles(document)
    _setup_a4_landscape(document)

    document.add_heading(title, level=1)

    total = 0
    subsection_index = 0
    for section_title, findings in groups:
        if not findings:
            continue

        subsection_index += 1
        heading_text = f"{section_base_number}.{subsection_index} {section_title or 'Findings'}"
        document.add_heading(heading_text, level=2)

        intro = document.add_paragraph("The following issues are identified:")
        intro.paragraph_format.space_before = Pt(0)
        intro.paragraph_format.space_after = Pt(6)
        for run in intro.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)

        add_section_summary_table(document, findings, hmap)
        total += len(findings)

        # Spacer between this section's table and the next section heading.
        spacer = document.add_paragraph("")
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

    if total == 0:
        warn("No findings were extracted - the output document will be empty of tables.")

    return document


# =============================================================================
# "veri-summary-by-section" format: one verification-status-count table per
# section, A4 portrait ("Verification Summary per Section")
# =============================================================================

# Base (always-shown) risk levels, in display order. "critical" is inserted
# at the front of this list at render time ONLY if used anywhere in the
# workbook (see compute_veri_summary_flags()).
_VERI_SUMMARY_BASE_RISK_LEVELS = ["high", "medium", "low", "ofi"]

# Base (always-shown) status columns, in display order. "Partially Completed"
# is inserted between "Completed" and "Incomplete" ONLY if used anywhere in
# the workbook (see compute_veri_summary_flags()).
_VERI_SUMMARY_BASE_STATUS_COLUMNS = ["Completed", "Incomplete", "Scheduled", "Accepted"]

# (label, width) for the two always-present leading columns.
VERI_SUMMARY_LEADING_COLUMNS = [
    ("Risk Level", Cm(2.6)),
    ("Total", Cm(1.6)),
]
# width used for EVERY status column (Completed/Partially Completed/
# Incomplete/Scheduled/Accepted) - kept uniform for a clean grid, wide
# enough that "Incomplete"/"Partially Completed" (the two longest labels)
# don't wrap awkwardly mid-word.
VERI_SUMMARY_STATUS_COL_WIDTH = Cm(2.6)


def _setup_a4_portrait(document: Document) -> None:
    """Configure the document's first section as A4, portrait orientation."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def compute_veri_summary_flags(groups: list[tuple[Optional[str], list[Finding]]]) -> tuple[bool, bool]:
    """Scan ALL findings across ALL sections (globally, not per-section) to
    decide once, for the WHOLE document:
        - has_critical: whether the "Critical" risk-level row should be
          shown (above "High") in every section's table.
        - has_partial: whether the "Partially Completed" status column
          should be shown (between "Completed" and "Incomplete") in every
          section's table.
    Deciding this globally (rather than per-section) keeps every section's
    table the same shape, which is the cleanest layout for a multi-section
    report."""
    has_critical = False
    has_partial = False
    for _, findings in groups:
        for f in findings:
            if normalize(f.risk_level) == "critical":
                has_critical = True
            if f.verification_status == "Partially Completed":
                has_partial = True
    return has_critical, has_partial


def _veri_summary_risk_levels(has_critical: bool) -> list[str]:
    return (["critical"] if has_critical else []) + _VERI_SUMMARY_BASE_RISK_LEVELS


def _veri_summary_status_columns(has_partial: bool) -> list[str]:
    cols = list(_VERI_SUMMARY_BASE_STATUS_COLUMNS)
    if has_partial:
        cols.insert(1, "Partially Completed")  # between Completed and Incomplete
    return cols


def _count_section_by_risk_and_status(
    findings: list[Finding], risk_levels: list[str], status_columns: list[str]
) -> tuple[dict, dict]:
    """Returns (counts, totals):
        - counts[risk_level][status_column] = number of findings in this
          section with that risk level AND that (canonical) verification
          status.
        - totals[risk_level] = total number of findings in this section
          with that risk level (regardless of verification status,
          including unmatched/invalid statuses - those are already flagged
          via warnings elsewhere and simply don't add to any status column
          here, but DO still count towards the risk level's Total)."""
    counts = {rl: {col: 0 for col in status_columns} for rl in risk_levels}
    totals = {rl: 0 for rl in risk_levels}

    for f in findings:
        rl_norm = normalize(f.risk_level)
        if rl_norm not in risk_levels:
            # Either an unrecognized risk level (already warned elsewhere)
            # or - for "critical" - a level not shown because has_critical
            # was globally False (shouldn't happen, since has_critical is
            # computed FROM these same findings, but guarded defensively).
            continue
        totals[rl_norm] += 1
        if f.verification_status in status_columns:
            counts[rl_norm][f.verification_status] += 1

    return counts, totals


def add_veri_summary_table(
    document: Document,
    section_title: str,
    findings: list[Finding],
    risk_levels: list[str],
    status_columns: list[str],
) -> None:
    """Render ONE section's "Verification Summary" table: counts of
    findings by Risk Level x Rectification (verification) Status."""
    counts, totals = _count_section_by_risk_and_status(findings, risk_levels, status_columns)

    n_status_cols = len(status_columns)
    n_cols = 2 + n_status_cols  # Risk Level, Total, + status columns
    column_widths = [w for _, w in VERI_SUMMARY_LEADING_COLUMNS] + [VERI_SUMMARY_STATUS_COL_WIDTH] * n_status_cols

    n_rows = 3 + len(risk_levels) + 1  # title + 2 header rows + risk rows + total row
    table = document.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_column_widths(table, column_widths)

    # ---- Row 0: full-width section title, blue/white/bold ----
    title_row = table.rows[0]
    title_cell = title_row.cells[0]
    for c in title_row.cells[1:]:
        title_cell = title_cell.merge(c)
    _set_cell_text(title_cell, [f"Security Risk Assessment - {section_title}"], bold=True, font_color=HEADER_ROW_FONT_COLOR)
    _set_cell_fill(title_cell, HEADER_ROW_FILL)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ---- Rows 1-2: header block ----
    header_row_a = table.rows[1]
    header_row_b = table.rows[2]

    # "Risk Level" - merged vertically across both header rows.
    risk_level_header_cell = header_row_a.cells[0].merge(header_row_b.cells[0])
    _set_cell_text(risk_level_header_cell, ["Risk Level"], bold=True, font_color=HEADER_ROW_FONT_COLOR)
    _set_cell_fill(risk_level_header_cell, HEADER_ROW_FILL)
    risk_level_header_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Column 1 (blank in row A, "Total" in row B) - NOT merged vertically,
    # matching the reference sample's layout.
    _set_cell_text(header_row_a.cells[1], [""], bold=True, font_color=HEADER_ROW_FONT_COLOR)
    _set_cell_fill(header_row_a.cells[1], HEADER_ROW_FILL)
    _set_cell_text(header_row_b.cells[1], ["Total"], bold=True, font_color=HEADER_ROW_FONT_COLOR)
    _set_cell_fill(header_row_b.cells[1], HEADER_ROW_FILL)

    # "Number of items by Rectification Status" - merged horizontally across
    # all status columns, in row A only.
    super_header_cell = header_row_a.cells[2]
    for c in header_row_a.cells[3:]:
        super_header_cell = super_header_cell.merge(c)
    _set_cell_text(super_header_cell, ["Number of items by Rectification Status"], bold=True, font_color=HEADER_ROW_FONT_COLOR)
    _set_cell_fill(super_header_cell, HEADER_ROW_FILL)

    # Individual status column names in row B.
    for i, col_name in enumerate(status_columns):
        cell = header_row_b.cells[2 + i]
        _set_cell_text(cell, [col_name], bold=True, font_color=HEADER_ROW_FONT_COLOR)
        _set_cell_fill(cell, HEADER_ROW_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for row in (title_row, header_row_a, header_row_b):
        _set_repeat_header_row(row)

    # ---- Risk-level data rows ----
    for i, rl in enumerate(risk_levels):
        row = table.rows[3 + i]
        label_cell = row.cells[0]
        _set_cell_text(label_cell, [RISK_LEVEL_DISPLAY[rl]], bold=True)
        # Reuse the EXACT SAME risk-level color scheme as the other formats.
        _set_cell_fill(label_cell, RISK_LEVEL_COLORS.get(rl))
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        total_cell = row.cells[1]
        _set_cell_text(total_cell, [str(totals[rl])], bold=False)
        _set_cell_fill(total_cell, None)
        total_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        for j, col in enumerate(status_columns):
            cell = row.cells[2 + j]
            _set_cell_text(cell, [str(counts[rl][col])], bold=False)
            _set_cell_fill(cell, None)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ---- Final "Total" row: label not bold, all counts bold ----
    total_row = table.rows[3 + len(risk_levels)]
    _set_cell_text(total_row.cells[0], ["Total"], bold=False)
    _set_cell_fill(total_row.cells[0], None)
    total_row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    grand_total = sum(totals.values())
    cell = total_row.cells[1]
    _set_cell_text(cell, [str(grand_total)], bold=True)
    _set_cell_fill(cell, None)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for j, col in enumerate(status_columns):
        col_total = sum(counts[rl][col] for rl in risk_levels)
        cell = total_row.cells[2 + j]
        _set_cell_text(cell, [str(col_total)], bold=True)
        _set_cell_fill(cell, None)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def build_veri_summary_document(
    groups: list[tuple[Optional[str], list[Finding]]],
    title: str,
) -> Document:
    """Build the "veri-summary-by-section" output: A4 portrait, one compact
    Risk-Level x Rectification-Status count table PER SECTION, reusing the
    same Finding data produced by extract_findings()."""
    document = Document()
    _apply_base_styles(document)
    _setup_a4_portrait(document)

    document.add_heading(title, level=1)

    has_critical, has_partial = compute_veri_summary_flags(groups)
    risk_levels = _veri_summary_risk_levels(has_critical)
    status_columns = _veri_summary_status_columns(has_partial)

    total = 0
    for section_title, findings in groups:
        if not findings:
            continue
        add_veri_summary_table(document, section_title or "Findings", findings, risk_levels, status_columns)
        total += len(findings)
        _add_blank_separator_paragraph(document)
        _add_blank_separator_paragraph(document)

    if total == 0:
        warn("No findings were extracted - the output document will be empty of tables.")

    return document


# =============================================================================
# Main
# =============================================================================


def convert(
    input_path: Path,
    output_path: Path,
    sheet_name: str = DEFAULT_SHEET_NAME,
    top_row: Optional[int] = None,
    left_col: Optional[int] = None,
    bottom_row: Optional[int] = None,
    right_col: Optional[int] = None,
    output_format: str = DEFAULT_OUTPUT_FORMAT,
    section_number: str = DEFAULT_SECTION_NUMBER,
    debug: bool = False,
) -> None:
    wb = openpyxl.load_workbook(input_path, data_only=True)
    ws = load_sheet(wb, sheet_name)

    if None not in (top_row, left_col, bottom_row, right_col):
        min_row, min_col, max_row, max_col = top_row, left_col, bottom_row, right_col
    else:
        min_row, min_col, max_row, max_col = find_table_bounds(ws)

    if debug:
        print(
            f"[debug] Sheet: {ws.title!r} | Table bounds: "
            f"row {min_row}-{max_row}, col {min_col}-{max_col}",
            file=sys.stderr,
        )

    hmap = map_headers(ws, header_row=min_row, min_col=min_col, max_col=max_col)

    if debug:
        print(f"[debug] Header map: {hmap}", file=sys.stderr)

    groups = extract_findings(ws, hmap, header_row=min_row, min_col=min_col, max_col=max_col, max_row=max_row)

    if debug:
        n = sum(len(f) for _, f in groups)
        print(f"[debug] Extracted {n} findings across {len(groups)} section group(s).", file=sys.stderr)

    if output_format == FORMAT_LANDSCAPE_DETAIL:
        document = build_landscape_document(
            groups, hmap, title="Follow-up Findings", section_base_number=section_number
        )
    elif output_format == FORMAT_VERI_SUMMARY_BY_SECTION:
        document = build_veri_summary_document(groups, title="Verification Summary by Section")
    else:
        document = build_document(groups, title="Follow-up Findings")
    document.save(output_path)

    print(f"Saved: {output_path}")
    if WARNINGS:
        print(f"\n{len(WARNINGS)} warning(s) were raised during conversion:", file=sys.stderr)
        for w in WARNINGS:
            print(f"  - {w}", file=sys.stderr)
    else:
        print("No validation warnings.", file=sys.stderr)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert a cybersecurity follow-up plan Excel workbook into a Word report of findings."
    )
    parser.add_argument("input", type=Path, help="Path to the source .xlsx workbook")
    parser.add_argument("output", type=Path, nargs="?", default=None, help="Path to the output .docx file (default: <input>.docx)")
    parser.add_argument("--sheet", default=DEFAULT_SHEET_NAME, help='Worksheet name to read (default: "Follow-up Items")')
    parser.add_argument("--top-row", type=int, default=None, help="Manually override: 1-based header row")
    parser.add_argument("--left-col", type=int, default=None, help="Manually override: 1-based leftmost column")
    parser.add_argument("--bottom-row", type=int, default=None, help="Manually override: 1-based last data row")
    parser.add_argument("--right-col", type=int, default=None, help="Manually override: 1-based rightmost column")
    parser.add_argument(
        "--format",
        dest="output_format",
        choices=OUTPUT_FORMATS,
        default=DEFAULT_OUTPUT_FORMAT,
        help=(
            'Output format (default: "portrait-detail"). '
            '"portrait-detail" = A4 portrait, one detailed table per finding. '
            '"landscape-detail" = A4 landscape, one summary table per section. '
            '"veri-summary-by-section" = A4 portrait, one verification-status-count table per section.'
        ),
    )
    parser.add_argument(
        "--section-number",
        default=DEFAULT_SECTION_NUMBER,
        help=(
            'Base report section number used to auto-number section headings '
            'in "landscape-detail" (default: "9", producing "9.1", "9.2", ...). '
            'Ignored for "portrait-detail" and "veri-summary-by-section".'
        ),
    )
    parser.add_argument("--debug", action="store_true", help="Print diagnostic information while converting")
    args = parser.parse_args()

    output = args.output or args.input.with_suffix(".docx")

    convert(
        input_path=args.input,
        output_path=output,
        sheet_name=args.sheet,
        top_row=args.top_row,
        left_col=args.left_col,
        bottom_row=args.bottom_row,
        right_col=args.right_col,
        output_format=args.output_format,
        section_number=args.section_number,
        debug=args.debug,
    )


if __name__ == "__main__":
    main()
